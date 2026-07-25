from __future__ import annotations

import io
from xml.sax.saxutils import escape

from .core import AnalysisResult


def build_docx_report(result: AnalysisResult, document_name: str) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    heading_tokens = {
        "Title": (24, "102A43", 0, 6),
        "Heading 1": (16, "1769AA", 14, 7),
        "Heading 2": (13, "1769AA", 10, 5),
        "Heading 3": (11.5, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "明白金 FinFair｜金融产品公平说明书"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(91, 107, 125)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("教学模拟｜第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    footer.add_run(" 页")

    doc.add_heading("金融产品公平说明书", level=0)
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"分析文件：{result.document_name or document_name}\n").bold = True
    subtitle.add_run(f"分析引擎：{result.engine}\n")
    subtitle.add_run(f"分析模式：{result.analysis_mode}\n")
    subtitle.add_run(
        f"文档覆盖：解析{result.page_count}页，提取{result.extracted_char_count}个字符，"
        f"空白/未提取文字页面{result.empty_page_count}页\n"
    )
    if result.agent_run.enabled:
        subtitle.add_run(
            f"Agent覆盖：接收{result.agent_char_count}个文档字符；"
            f"{'输入已截断，未覆盖完整提取文本' if result.agent_truncated else '输入未截断'}\n"
        )
    else:
        subtitle.add_run("Agent覆盖：未启用，确定性规则处理全部已提取文字\n")
    subtitle.add_run("用途：教学演示，不构成投资建议或正式适当性评估。")

    doc.add_heading("30秒看懂产品", level=1)
    for item in result.fields:
        doc.add_heading(f"{item.label}：{item.value}", level=2)
        doc.add_paragraph(item.plain_language)
        page = f"第{item.evidence.page}页" if item.evidence.page else "未定位"
        evidence = doc.add_paragraph()
        evidence.add_run(f"证据（{page}）：").bold = True
        evidence.add_run(item.evidence.text)

    doc.add_heading("主要风险", level=1)
    for item in result.risks:
        page = f"第{item.evidence.page}页" if item.evidence.page else "未定位"
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{item.label}：{item.value}（{page}）").bold = True
        doc.add_paragraph(item.plain_language)

    status_labels = {
        "supported": "支持",
        "omitted": "未披露",
        "weakened": "弱化",
        "conflicting": "冲突",
        "unclear": "无法判断",
    }
    doc.add_heading("宣传材料与正式说明书逐项对照", level=1)
    if not result.marketing_provided:
        doc.add_paragraph("本次未输入宣传材料，因此没有执行宣传材料对照。")
    elif result.findings:
        for finding in result.findings:
            doc.add_heading(
                f"{finding.title}｜{status_labels.get(finding.status, finding.status)}",
                level=2,
            )
            doc.add_paragraph(finding.explanation)
            marketing = doc.add_paragraph()
            marketing.add_run("宣传材料说法：").bold = True
            marketing.add_run(finding.marketing_text or "未披露")
            formal = doc.add_paragraph()
            formal.add_run("正式说明书怎么写：").bold = True
            formal.add_run(finding.formal_plain_language)
            status = doc.add_paragraph()
            status.add_run("核验状态：").bold = True
            status.add_run(status_labels.get(finding.status, finding.status))
            page = (
                f"第{finding.formal_evidence.page}页"
                if finding.formal_evidence.page
                else "未定位"
            )
            evidence = doc.add_paragraph()
            evidence.add_run(f"正式文件证据（{page}）：").bold = True
            evidence.add_run(finding.formal_evidence.text)
    else:
        doc.add_paragraph("已输入宣传材料，当前规则没有生成可对照项目。")

    doc.add_heading("大模型语义增强", level=1)
    if result.agent_run.enabled:
        doc.add_paragraph(
            f"模型：{result.agent_run.model}；协议："
            f"{result.agent_run.protocol or '未记录'}。"
        )
        doc.add_paragraph(
            "固定流程：规则引擎 → 分析 Agent → 核验 Agent → 程序逐字引用门控。"
        )
        doc.add_paragraph(
            f"候选{result.agent_run.candidate_count}项；"
            f"核验支持{result.agent_run.verifier_supported_count}项；"
            f"门控通过{result.agent_run.gate_passed_count}项；"
            f"最终拦截{result.agent_run.rejected_count}项。"
        )
        doc.add_paragraph(f"停止条件：{result.agent_run.stop_reason}")
        if result.agent_run.error:
            doc.add_paragraph(f"降级原因：{result.agent_run.error}")
        for reason in result.agent_run.rejection_reasons:
            doc.add_paragraph(f"拦截：{reason}", style="List Bullet")
    if result.agent_run.enabled and result.agent_insights:
        for insight in result.agent_insights:
            doc.add_heading(insight.title, level=2)
            doc.add_paragraph(insight.conclusion)
            doc.add_paragraph(insight.plain_language)
            page = f"第{insight.evidence.page}页" if insight.evidence.page else "未定位"
            evidence = doc.add_paragraph()
            evidence.add_run(f"{insight.verification_status}｜原文（{page}）：").bold = True
            evidence.add_run(insight.evidence.text)
    elif result.agent_run.enabled:
        doc.add_paragraph("大模型增强已运行，但没有产生通过双重证据核验的新洞察。")
    else:
        doc.add_paragraph("本次使用规则模式，未启用大模型语义增强。")

    doc.add_heading("购买前必须确认的问题", level=1)
    for question in result.questions:
        doc.add_paragraph(question, style="List Number")

    doc.add_heading("局限与免责声明", level=1)
    for item in result.limitations:
        doc.add_paragraph(item, style="List Bullet")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_pdf_report(result: AnalysisResult, document_name: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        KeepTogether,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.72 * inch,
        leftMargin=0.72 * inch,
        topMargin=0.68 * inch,
        bottomMargin=0.62 * inch,
        title="金融产品公平说明书",
        author="FinFair",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CNBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor("#243B53"),
        spaceAfter=5,
    )
    title = ParagraphStyle(
        "CNTitle",
        parent=body,
        fontSize=22,
        leading=28,
        textColor=colors.HexColor("#102A43"),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h1 = ParagraphStyle(
        "CNH1",
        parent=body,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#1769AA"),
        spaceBefore=12,
        spaceAfter=7,
        keepWithNext=True,
    )
    h2 = ParagraphStyle(
        "CNH2",
        parent=body,
        fontSize=11.5,
        leading=16,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True,
    )
    evidence_style = ParagraphStyle(
        "Evidence",
        parent=body,
        fontSize=8.8,
        leading=13,
        leftIndent=10,
        borderColor=colors.HexColor("#1769AA"),
        borderWidth=0,
        borderPadding=7,
        backColor=colors.HexColor("#EDF5FB"),
        spaceAfter=8,
    )
    meta = ParagraphStyle(
        "Meta",
        parent=body,
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor("#526D82"),
        alignment=TA_CENTER,
        spaceAfter=3,
    )

    story = [Paragraph("金融产品公平说明书", title)]
    story.extend(
        [
            Paragraph(f"分析文件：{escape(result.document_name or document_name)}", meta),
            Paragraph(f"分析引擎：{escape(result.engine)}", meta),
            Paragraph(f"分析模式：{escape(result.analysis_mode)}", meta),
            Paragraph(
                f"文档覆盖：解析{result.page_count}页，"
                f"提取{result.extracted_char_count}个字符，"
                f"空白/未提取文字页面{result.empty_page_count}页",
                meta,
            ),
            Paragraph(
                (
                    f"Agent覆盖：接收{result.agent_char_count}个文档字符；"
                    f"{'输入已截断，未覆盖完整提取文本' if result.agent_truncated else '输入未截断'}"
                    if result.agent_run.enabled
                    else "Agent覆盖：未启用，确定性规则处理全部已提取文字"
                ),
                meta,
            ),
            Paragraph("教学演示，不构成投资建议或正式适当性评估。", meta),
            Spacer(1, 8),
            Paragraph("30秒看懂产品", h1),
        ]
    )
    for item in result.fields:
        page = f"第{item.evidence.page}页" if item.evidence.page else "未定位"
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f"{escape(item.label)}：{escape(item.value)}", h2
                    ),
                    Paragraph(escape(item.plain_language), body),
                    Paragraph(
                        f"<b>证据（{page}）：</b>{escape(item.evidence.text)}",
                        evidence_style,
                    ),
                ]
            )
        )

    story.append(Paragraph("主要风险", h1))
    for item in result.risks:
        page = f"第{item.evidence.page}页" if item.evidence.page else "未定位"
        story.append(
            Paragraph(
                f"• <b>{escape(item.label)}：</b>{escape(item.value)}（{page}）",
                body,
            )
        )
        story.append(Paragraph(escape(item.plain_language), evidence_style))

    status_labels = {
        "supported": "支持",
        "omitted": "未披露",
        "weakened": "弱化",
        "conflicting": "冲突",
        "unclear": "无法判断",
    }
    story.append(Paragraph("宣传材料与正式说明书逐项对照", h1))
    if not result.marketing_provided:
        story.append(Paragraph("本次未输入宣传材料，因此没有执行宣传材料对照。", body))
    elif result.findings:
        for finding in result.findings:
            page = (
                f"第{finding.formal_evidence.page}页"
                if finding.formal_evidence.page
                else "未定位"
            )
            story.append(
                KeepTogether(
                    [
                        Paragraph(
                            f"{escape(finding.title)}｜"
                            f"{escape(status_labels.get(finding.status, finding.status))}",
                            h2,
                        ),
                        Paragraph(escape(finding.explanation), body),
                        Paragraph(
                            f"<b>宣传材料说法：</b>{escape(finding.marketing_text or '未披露')}",
                            body,
                        ),
                        Paragraph(
                            f"<b>正式说明书怎么写：</b>{escape(finding.formal_plain_language)}",
                            body,
                        ),
                        Paragraph(
                            f"<b>核验状态：</b>"
                            f"{escape(status_labels.get(finding.status, finding.status))}",
                            body,
                        ),
                        Paragraph(
                            f"<b>正式证据（{page}）：</b>{escape(finding.formal_evidence.text)}",
                            evidence_style,
                        ),
                    ]
                )
            )
    else:
        story.append(Paragraph("已输入宣传材料，当前规则没有生成可对照项目。", body))

    story.append(Paragraph("大模型语义增强", h1))
    if result.agent_run.enabled:
        story.extend(
            [
                Paragraph(
                    f"模型：{escape(result.agent_run.model)}；协议："
                    f"{escape(result.agent_run.protocol or '未记录')}。",
                    body,
                ),
                Paragraph(
                    "固定流程：规则引擎 → 分析 Agent → 核验 Agent → 程序逐字引用门控。",
                    body,
                ),
                Paragraph(
                    f"候选{result.agent_run.candidate_count}项；"
                    f"核验支持{result.agent_run.verifier_supported_count}项；"
                    f"门控通过{result.agent_run.gate_passed_count}项；"
                    f"最终拦截{result.agent_run.rejected_count}项。",
                    body,
                ),
                Paragraph(
                    f"停止条件：{escape(result.agent_run.stop_reason)}", body
                ),
            ]
        )
        if result.agent_run.error:
            story.append(
                Paragraph(f"降级原因：{escape(result.agent_run.error)}", body)
            )
        for reason in result.agent_run.rejection_reasons:
            story.append(Paragraph(f"• 拦截：{escape(reason)}", body))
    if result.agent_run.enabled and result.agent_insights:
        for insight in result.agent_insights:
            page = f"第{insight.evidence.page}页" if insight.evidence.page else "未定位"
            story.append(
                KeepTogether(
                    [
                        Paragraph(escape(insight.title), h2),
                        Paragraph(escape(insight.conclusion), body),
                        Paragraph(escape(insight.plain_language), body),
                        Paragraph(
                            f"<b>{escape(insight.verification_status)}｜原文（{page}）：</b>"
                            f"{escape(insight.evidence.text)}",
                            evidence_style,
                        ),
                    ]
                )
            )
    elif result.agent_run.enabled:
        story.append(Paragraph("没有产生通过双重证据核验的新洞察。", body))
    else:
        story.append(Paragraph("本次使用规则模式，未启用大模型语义增强。", body))

    question_style = ParagraphStyle(
        "Question",
        parent=body,
        fontSize=9,
        leading=12,
        spaceAfter=1,
    )
    questions = [
        ListItem(Paragraph(escape(question), question_style), leftIndent=12)
        for question in result.questions
    ]
    closing_block = [
        Paragraph("购买前必须确认的问题", h1),
        ListFlowable(questions, bulletType="1", leftIndent=22),
        Paragraph("局限与免责声明", h1),
    ]
    for item in result.limitations:
        closing_block.append(Paragraph(f"• {escape(item)}", body))
    story.append(KeepTogether(closing_block))

    def add_page_number(canvas, pdf_doc):
        canvas.saveState()
        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor(colors.HexColor("#6B7C93"))
        canvas.drawString(0.72 * inch, 0.35 * inch, "明白金 FinFair｜教学模拟")
        canvas.drawRightString(
            letter[0] - 0.72 * inch,
            0.35 * inch,
            f"第 {pdf_doc.page} 页",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return buffer.getvalue()
